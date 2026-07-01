# Licensed to the Apache Software Foundation (ASF) under one
# or more contributor license agreements.  See the NOTICE file
# distributed with this work for additional information
# regarding copyright ownership.  The ASF licenses this file
# to you under the Apache License, Version 2.0 (the
# "License"); you may not use this file except in compliance
# with the License.  You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Unified Document Parser for HugeGraph-AI — handles DOCX, Markdown, PDF, and plain text.

Design principles (borrowed from LightRAG):
  - Template method pattern: every sub-parser returns the SAME DocumentExtractionResult
  - Heading-level block splitting (LightRAG docx/parse_document.py style)
  - Table extraction with merged-cell awareness (LightRAG table_extractor.py style)
  - Image relationship resolution (LightRAG drawing_image_extractor.py style)
  - Markdown regex-based extraction (LightRAG parser/markdown/extract.py style)

Operator protocol::

    parser = UnifiedDocumentParser()
    context = parser.run({
        "document_path": "/path/to/file.pdf",   # or "document_content": bytes
    })
    # context["document_extraction"] → DocumentExtractionResult
"""

from __future__ import annotations

import base64
import io
import json
import logging
import re
import uuid
import zipfile
from dataclasses import dataclass, field
from pathlib import Path, PurePosixPath
from typing import Any, Callable, Dict, List, Optional, Tuple, Union

log = logging.getLogger(__name__)

# ============================================================================
# Unified output dataclasses
# ============================================================================


@dataclass
class ContentBlock:
    """A single content block extracted from any document format.

    All sub-parsers produce ContentBlock instances so downstream consumers
    (chunk splitter, sidecar writer, VLM pipeline) can operate format-agnostic.
    """

    block_id: str
    heading: str
    heading_level: int  # 0 = no heading, 1-6 = ATX levels
    content: str  # body text (may include inline placeholders)
    block_type: str  # "paragraph" / "table" / "image" / "equation"
    page_num: Optional[int] = None  # for PDF pages
    source_id: str = ""  # reference back to original position (paraId, etc.)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "block_id": self.block_id,
            "heading": self.heading,
            "heading_level": self.heading_level,
            "content": self.content,
            "block_type": self.block_type,
            "page_num": self.page_num,
            "source_id": self.source_id,
        }


@dataclass
class ExtractedImage:
    """An image extracted from any document format.

    Stores base64-encoded data plus metadata for VLM consumption.
    """

    image_id: str
    base64_data: str
    format: str = "jpeg"  # original format (jpeg/png/gif/etc.)
    size: Tuple[int, int] = (0, 0)  # (width, height) pixels
    caption: str = ""
    page_num: Optional[int] = None  # PDF page number
    source_id: str = ""  # relationship id (DOCX) or reference key
    alt_text: str = ""  # markdown alt text

    @property
    def data_uri(self) -> str:
        mime = f"image/{self.format}" if self.format != "jpg" else "image/jpeg"
        return f"data:{mime};base64,{self.base64_data}"

    def to_dict(self) -> Dict[str, Any]:
        return {
            "image_id": self.image_id,
            "format": self.format,
            "size": self.size,
            "caption": self.caption,
            "page_num": self.page_num,
            "source_id": self.source_id,
            "alt_text": self.alt_text,
            "base64_data": self.base64_data,
        }


@dataclass
class ExtractedTable:
    """A table extracted from any document format.

    Stores rows as a 2D string array; header_rows is separated when available.
    """

    table_id: str
    rows: List[List[str]] = field(default_factory=list)
    header_rows: Optional[List[List[str]]] = None
    caption: str = ""
    num_rows: int = 0
    num_cols: int = 0
    page_num: Optional[int] = None
    source_id: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return {
            "table_id": self.table_id,
            "rows": self.rows,
            "header_rows": self.header_rows,
            "caption": self.caption,
            "num_rows": self.num_rows,
            "num_cols": self.num_cols,
            "page_num": self.page_num,
            "source_id": self.source_id,
        }


@dataclass
class ExtractedEquation:
    """An equation extracted from any document format.

    Stores LaTeX source; is_block distinguishes display vs inline math.
    """

    equation_id: str
    latex: str
    is_block: bool = True
    caption: str = ""
    page_num: Optional[int] = None
    source_id: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return {
            "equation_id": self.equation_id,
            "latex": self.latex,
            "is_block": self.is_block,
            "caption": self.caption,
            "page_num": self.page_num,
            "source_id": self.source_id,
        }


@dataclass
class DocumentExtractionResult:
    """Unified extraction result that all format parsers produce.

    This is the single output contract for the UnifiedDocumentParser operator.
    """

    source_file: str
    source_format: str  # "pdf" / "docx" / "md" / "txt"
    blocks: List[ContentBlock] = field(default_factory=list)
    images: List[ExtractedImage] = field(default_factory=list)
    tables: List[ExtractedTable] = field(default_factory=list)
    equations: List[ExtractedEquation] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    @property
    def total_blocks(self) -> int:
        return len(self.blocks)

    @property
    def total_images(self) -> int:
        return len(self.images)

    @property
    def total_tables(self) -> int:
        return len(self.tables)

    @property
    def total_equations(self) -> int:
        return len(self.equations)

    @property
    def total_text_length(self) -> int:
        return sum(len(b.content) for b in self.blocks)

    def summary(self) -> Dict[str, Any]:
        return {
            "source": self.source_file,
            "format": self.source_format,
            "blocks": self.total_blocks,
            "images": self.total_images,
            "tables": self.total_tables,
            "equations": self.total_equations,
            "total_chars": self.total_text_length,
            "metadata": self.metadata,
        }

    def to_dict(self) -> Dict[str, Any]:
        return {
            "source_file": self.source_file,
            "source_format": self.source_format,
            "blocks": [b.to_dict() for b in self.blocks],
            "images": [i.to_dict() for i in self.images],
            "tables": [t.to_dict() for t in self.tables],
            "equations": [e.to_dict() for e in self.equations],
            "metadata": self.metadata,
        }


# ============================================================================
# Parser Registry — extensible format → parser mapping
# ============================================================================


class ParserRegistry:
    """Registry of format suffix → parser function mapping.

    Extensible: call register() to add custom parsers for new formats.
    """

    _parsers: Dict[str, Callable] = {}
    _aliases: Dict[str, str] = {
        ".markdown": ".md",
        ".mkd": ".md",
        ".mkdn": ".md",
        ".mdown": ".md",
        ".doc": ".docx",  # Will fail if it's a legacy .doc; user must convert
        ".text": ".txt",
        ".plaintext": ".txt",
        ".log": ".txt",
    }

    @classmethod
    def register(cls, suffix: str, parser_fn: Callable) -> None:
        """Register a parser function for a given file suffix.

        Args:
            suffix: File extension including dot (e.g. ".epub").
            parser_fn: Callable that accepts (file_path_or_bytes) and returns
                DocumentExtractionResult.
        """
        normalized = suffix.lower()
        cls._parsers[normalized] = parser_fn
        log.info(f"Registered parser for format: {normalized}")

    @classmethod
    def get_parser(cls, suffix: str) -> Optional[Callable]:
        """Look up the parser function for a file suffix.

        Resolves aliases first (e.g. ".markdown" → ".md").
        Returns None when no parser is registered for the suffix.
        """
        normalized = suffix.lower()
        resolved = cls._aliases.get(normalized, normalized)
        return cls._parsers.get(resolved)

    @classmethod
    def supported_formats(cls) -> List[str]:
        """Return all registered format suffixes."""
        all_suffixes = set(cls._parsers.keys()) | set(cls._aliases.keys())
        return sorted(all_suffixes)


# ============================================================================
# PDF sub-parser — reuses PyMuPDF approach from pdf_image_extractor.py
# ============================================================================

_MAX_IMAGE_SIZE_KB = 512
_JPEG_QUALITY = 75
_MIN_IMAGE_DIM = 50
_MAX_IMAGES_PER_PAGE = 50


def _parse_pdf(source: Union[str, Path, bytes]) -> DocumentExtractionResult:
    """Parse PDF using PyMuPDF (fitz), adapting output to ContentBlock format.

    Reuses the approach from pdf_image_extractor.py: extract images via
    page.get_images(), extract text via page.get_text("dict"), then convert
    into the unified DocumentExtractionResult.

    Args:
        source: File path string/Path, or raw PDF bytes.

    Returns:
        DocumentExtractionResult with blocks, images, tables, equations.
    """
    import fitz  # PyMuPDF

    # --- open PDF from path or bytes ---
    if isinstance(source, bytes):
        doc = fitz.open(stream=source, filetype="pdf")
        source_label = "<bytes>"
    else:
        pdf_path = str(Path(source).resolve())
        doc = fitz.open(pdf_path)
        source_label = pdf_path

    total_pages = len(doc)
    log.info(f"PDF opened: {source_label}, pages={total_pages}")

    result = DocumentExtractionResult(
        source_file=source_label,
        source_format="pdf",
        metadata={"total_pages": total_pages},
    )

    seen_xrefs: set = set()
    block_idx = 0
    img_idx = 0

    for page_num in range(total_pages):
        page = doc.load_page(page_num)
        page_rect = page.rect

        # --- extract images ---
        try:
            image_list = page.get_images(full=True)
        except Exception as e:
            log.warning(f"Failed to get images from page {page_num}: {e}")
            image_list = []

        for img_info in image_list:
            if img_idx >= _MAX_IMAGES_PER_PAGE * total_pages:
                log.warning("Reached global max image limit")
                break

            xref = img_info[0]
            width = img_info[2]
            height = img_info[3]

            if xref in seen_xrefs:
                continue
            if width < _MIN_IMAGE_DIM or height < _MIN_IMAGE_DIM:
                continue

            try:
                img_data = doc.extract_image(xref)
                if img_data is None:
                    continue
                raw_bytes = img_data["image"]
                original_ext = img_data["ext"]

                compressed_bytes = _compress_image_bytes(raw_bytes)
                b64_data = base64.b64encode(compressed_bytes).decode("ascii")

                extracted_img = ExtractedImage(
                    image_id=f"img_pdf_{page_num}_{img_idx}",
                    base64_data=b64_data,
                    format="jpeg",
                    size=(width, height),
                    page_num=page_num,
                    source_id=str(xref),
                )
                result.images.append(extracted_img)
                seen_xrefs.add(xref)
                img_idx += 1
            except Exception as e:
                log.warning(f"Failed to extract image xref={xref} on page {page_num}: {e}")

        # --- extract text blocks ---
        try:
            page_dict = page.get_text("dict")
            blocks = page_dict.get("blocks", [])
        except Exception as e:
            log.warning(f"Failed to get text dict from page {page_num}: {e}")
            blocks = []

        for blk in blocks:
            if blk.get("type") != 0:
                continue

            bbox = blk.get("bbox", (0, 0, 0, 0))
            lines = blk.get("lines", [])
            text_parts: list[str] = []
            font_sizes: list[float] = []

            for line in lines:
                spans = line.get("spans", [])
                line_text = ""
                for span in spans:
                    line_text += span.get("text", "")
                    fs = span.get("size", 0)
                    if fs > 0:
                        font_sizes.append(fs)
                if line_text.strip():
                    text_parts.append(line_text.strip())

            text = "\n".join(text_parts).strip()
            if not text:
                continue

            # Heuristic heading detection: short text + larger font
            heading_level = 0
            heading = ""
            if font_sizes and len(text) < 100:
                avg_fs = sum(font_sizes) / len(font_sizes)
                if avg_fs > 14:
                    heading_level = 1
                    heading = text
                elif avg_fs > 12:
                    heading_level = 2
                    heading = text

            block_id = f"blk_pdf_{page_num}_{block_idx}"
            result.blocks.append(
                ContentBlock(
                    block_id=block_id,
                    heading=heading,
                    heading_level=heading_level,
                    content=text,
                    block_type="paragraph",
                    page_num=page_num,
                    source_id=block_id,
                )
            )
            block_idx += 1

    doc.close()
    log.info(f"PDF extraction complete: {result.summary()}")
    return result


def _compress_image_bytes(raw_bytes: bytes) -> bytes:
    """Compress image bytes to JPEG within size budget.

    Uses PIL for resize/quality reduction when the raw payload exceeds
    _MAX_IMAGE_SIZE_KB. Falls back to raw bytes when PIL is unavailable.
    """
    target_size = _MAX_IMAGE_SIZE_KB * 1024

    if len(raw_bytes) <= target_size:
        return raw_bytes

    try:
        from PIL import Image as PILImage

        img = PILImage.open(io.BytesIO(raw_bytes))
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")

        # Progressive quality reduction
        for quality in range(_JPEG_QUALITY, 20, -5):
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=quality, optimize=True)
            compressed = buf.getvalue()
            if len(compressed) <= target_size:
                return compressed

        # Scale reduction
        scale = 0.8
        while scale > 0.3:
            new_w = int(img.width * scale)
            new_h = int(img.height * scale)
            resized = img.resize((new_w, new_h), PILImage.LANCZOS)
            buf = io.BytesIO()
            resized.save(buf, format="JPEG", quality=_JPEG_QUALITY // 2, optimize=True)
            compressed = buf.getvalue()
            if len(compressed) <= target_size:
                return compressed
            scale -= 0.1

        # Final fallback: tiny image
        small = img.resize((img.width // 4, img.height // 4), PILImage.LANCZOS)
        buf = io.BytesIO()
        small.save(buf, format="JPEG", quality=30, optimize=True)
        return buf.getvalue()

    except Exception as e:
        log.error(f"Image compression failed: {e}")
        return raw_bytes


# ============================================================================
# DOCX sub-parser — uses python-docx with heading hierarchy, tables, images
# ============================================================================

# OOXML namespace shortcuts
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_V_NS = "urn:schemas-microsoft-com:vml"
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_IMAGE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"

_DOCX_NS = {
    "w": _W_NS,
    "wp": _WP_NS,
    "a": _A_NS,
    "r": _R_NS,
    "m": _M_NS,
}

# Revision / comment tags whose subtree must be dropped
_SKIP_REVISION_TAGS = frozenset({"del", "moveFrom"})
_SKIP_COMMENT_TAGS = frozenset(
    {"commentRangeStart", "commentRangeEnd", "commentReference", "annotationRef"}
)
_SKIP_PARA_TAGS = _SKIP_REVISION_TAGS | _SKIP_COMMENT_TAGS

MAX_HEADING_LENGTH = 200


@dataclass
class _DocxRelationship:
    """Minimal relationship metadata for image resolution."""

    rel_id: str
    target: str
    target_mode: str  # "Internal" / "External"
    rel_type: str
    part_name: Optional[str] = None
    content_type: Optional[str] = None
    image_format: Optional[str] = None


def _normalize_image_format(ext_or_type: str) -> Optional[str]:
    if not ext_or_type:
        return None
    value = ext_or_type.strip().lower()
    if value.startswith("image/"):
        value = value.split("/", 1)[1]
        if "+" in value:
            value = value.split("+", 1)[0]
        if value.startswith("x-"):
            value = value[2:]
    value = value.lstrip(".")
    if value == "jpg":
        return "jpeg"
    if value in {"jpeg", "png", "gif", "bmp", "tiff", "webp", "svg", "emf", "wmf"}:
        return value
    return value or None


def _resolve_part_name(source_part: str, target: str) -> str:
    import posixpath

    if target.startswith("/"):
        return posixpath.normpath(target)
    source_dir = posixpath.dirname(source_part)
    joined = posixpath.join(source_dir, target)
    normalized = posixpath.normpath(joined)
    if not normalized.startswith("/"):
        normalized = "/" + normalized
    return normalized


def _load_docx_relationships(docx_path: str) -> Dict[str, _DocxRelationship]:
    """Parse word/_rels/document.xml.rels + [Content_Types].xml from a DOCX zip."""
    from xml.etree import ElementTree as ET

    rels: Dict[str, _DocxRelationship] = {}
    source_part = "/word/document.xml"
    overrides: Dict[str, str] = {}
    defaults: Dict[str, str] = {}

    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            names = zf.namelist()

            # Content types
            if "[Content_Types].xml" in names:
                ct_root = ET.parse(zf.open("[Content_Types].xml")).getroot()
                for node in ct_root.findall(f".//{{{_CT_NS}}}Override"):
                    pn = node.get("PartName")
                    ct = node.get("ContentType")
                    if pn and ct:
                        overrides[pn] = ct
                for node in ct_root.findall(f".//{{{_CT_NS}}}Default"):
                    ext = node.get("Extension")
                    ct = node.get("ContentType")
                    if ext and ct:
                        defaults[ext.lower()] = ct

            # Relationships
            rels_xml = "word/_rels/document.xml.rels"
            if rels_xml not in names:
                return rels
            rels_root = ET.parse(zf.open(rels_xml)).getroot()

            for rel in rels_root.findall(f".//{{{_REL_NS}}}Relationship"):
                rel_id = rel.get("Id")
                target = rel.get("Target", "")
                target_mode = rel.get("TargetMode", "Internal")
                rel_type = rel.get("Type", "")
                if not rel_id:
                    continue

                part_name = None
                content_type = None
                image_format = None

                if target_mode.lower() != "external":
                    part_name = _resolve_part_name(source_part, target)
                    if part_name:
                        content_type = overrides.get(part_name)
                        if not content_type:
                            ext = PurePosixPath(part_name).suffix.lower().lstrip(".")
                            content_type = defaults.get(ext)
                        image_format = _normalize_image_format(content_type or "")
                else:
                    parsed_suffix = PurePosixPath(target).suffix
                    image_format = _normalize_image_format(parsed_suffix)

                rels[rel_id] = _DocxRelationship(
                    rel_id=rel_id,
                    target=target,
                    target_mode=target_mode,
                    rel_type=rel_type,
                    part_name=part_name,
                    content_type=content_type,
                    image_format=image_format,
                )
    except Exception as e:
        log.warning(f"Failed to load DOCX relationships: {e}")

    return rels


def _extract_embedded_image_bytes(docx_path: str, part_name: str) -> Optional[bytes]:
    """Read embedded image bytes from the DOCX zip archive."""
    zip_member = part_name.lstrip("/")
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            return zf.read(zip_member)
    except Exception:
        return None


def _parse_styles_outline_levels(docx_path: str) -> Dict[str, int]:
    """Parse styles.xml to extract outlineLvl definitions with inheritance."""
    from xml.etree import ElementTree as ET

    styles_outline: Dict[str, int] = {}
    style_based_on: Dict[str, str] = {}

    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            if "word/styles.xml" not in zf.namelist():
                return styles_outline
            tree = ET.parse(zf.open("word/styles.xml"))
            root = tree.getroot()

            for style_el in root.findall(f".//{{{_W_NS}}}style"):
                style_id = style_el.get(f"{{{_W_NS}}}styleId")
                if not style_id:
                    continue

                based_on = style_el.find(f"{{{_W_NS}}}basedOn")
                if based_on is not None:
                    parent_id = based_on.get(f"{{{_W_NS}}}val")
                    if parent_id:
                        style_based_on[style_id] = parent_id

                pPr = style_el.find(f"{{{_W_NS}}}pPr")
                if pPr is not None:
                    outline_el = pPr.find(f"{{{_W_NS}}}outlineLvl")
                    if outline_el is not None:
                        level = int(outline_el.get(f"{{{_W_NS}}}val"))
                        styles_outline[style_id] = level

        # Resolve inheritance
        def _get_level(sid: str, visited: set = None) -> Optional[int]:
            if visited is None:
                visited = set()
            if sid in visited:
                return None
            visited.add(sid)
            if sid in styles_outline:
                return styles_outline[sid]
            if sid in style_based_on:
                return _get_level(style_based_on[sid], visited)
            return None

        for sid in set(styles_outline.keys()) | set(style_based_on.keys()):
            if sid not in styles_outline:
                lvl = _get_level(sid)
                if lvl is not None:
                    styles_outline[sid] = lvl
    except Exception:
        pass

    return styles_outline


def _get_heading_level(para_element, styles_outline_map: Dict[str, int]) -> Optional[int]:
    """Get heading level from a paragraph element.

    Priority: paragraph outlineLvl > style outlineLvl.
    Returns 0-based level (0=H1) or None for non-heading paragraphs.
    """
    pPr = para_element.find(f"{{{_W_NS}}}pPr")
    if pPr is not None:
        outline_el = pPr.find(f"{{{_W_NS}}}outlineLvl")
        if outline_el is not None:
            level = int(outline_el.get(f"{{{_W_NS}}}val"))
            return level if level < 9 else None

    if pPr is not None:
        pStyle_el = pPr.find(f"{{{_W_NS}}}pStyle")
        if pStyle_el is not None:
            style_id = pStyle_el.get(f"{{{_W_NS}}}val")
            if style_id and style_id in styles_outline_map:
                level = styles_outline_map[style_id]
                return level if level < 9 else None

    # Fallback: check python-docx paragraph style name
    try:
        style_name = para_element.get("style_name", "")
        heading_match = re.match(r"^Heading\s*(\d)", style_name)
        if heading_match:
            return int(heading_match.group(1)) - 1  # 0-based
    except Exception:
        pass

    return None


def _extract_para_id(para_element) -> Optional[str]:
    """Extract w14:paraId from a paragraph element."""
    return para_element.get(
        "{http://schemas.microsoft.com/office/word/2010/wordml}paraId"
    )


def _extract_paragraph_content_omml(element, ns: Dict[str, str]) -> str:
    """Extract text + OMML equations from a paragraph element.

    Handles w:r (text runs), m:oMath (inline equations), m:oMathPara
    (block equations). Recurses into container elements.
    OMML → LaTeX conversion is simplified (basic elements only).
    """
    parts: list[str] = []

    def _append_from(node) -> None:
        tag = node.tag.split("}")[-1]
        if tag in _SKIP_PARA_TAGS:
            return
        if tag == "r":
            parts.append(_extract_run_text(node))
            return
        if tag == "oMath" or tag == "oMathPara":
            latex = _convert_omml_to_latex_simplified(node)
            if latex:
                parts.append(f"<equation>{latex}</equation>")
            return
        for child in node:
            _append_from(child)

    for child in element:
        _append_from(child)
    return "".join(parts)


def _extract_run_text(run_element) -> str:
    """Extract text from a w:r element, preserving basic formatting."""
    text = ""

    # Check superscript/subscript
    rPr = run_element.find(f"{{{_W_NS}}}rPr")
    vert_align = None
    if rPr is not None:
        vert_el = rPr.find(f"{{{_W_NS}}}vertAlign")
        if vert_el is not None:
            vert_align = vert_el.get(f"{{{_W_NS}}}val")

    for child in run_element:
        tag = child.tag.split("}")[-1]
        if tag == "t" and child.text:
            text += child.text
        elif tag == "tab":
            text += "\t"
        elif tag == "br":
            text += "\n"

    if text and vert_align == "superscript":
        return f"<sup>{text}</sup>"
    elif text and vert_align == "subscript":
        return f"<sub>{text}</sub>"
    return text


def _convert_omml_to_latex_simplified(element) -> str:
    """Simplified OMML → LaTeX conversion for basic equation elements.

    Handles: m:r (text runs), m:f (fractions), m:sup (superscripts),
    m:sub (subscripts), m:nary (n-ary operators like integrals),
    m:d (delimiters/brackets).

    This is intentionally simplified — full OMML conversion requires a
    dedicated parser (see lightrag/parser/docx/omml/). For production
    use, that module should be imported; this fallback handles common cases.
    """
    try:
        from lightrag.parser.docx.omml import convert_omml_to_latex

        return convert_omml_to_latex(element)
    except ImportError:
        pass

    # Simplified fallback
    tag = element.tag.split("}")[-1]

    if tag == "oMathPara":
        parts = []
        for child in element:
            child_tag = child.tag.split("}")[-1]
            if child_tag == "oMath":
                parts.append(_convert_omml_to_latex_simplified(child))
        return " ".join(parts) if parts else ""

    if tag != "oMath":
        return ""

    parts: list[str] = []
    for child in element:
        child_tag = child.tag.split("}")[-1]

        if child_tag == "r":
            # m:r → text run
            t_el = child.find(f"{{{_M_NS}}}t")
            if t_el is not None and t_el.text:
                parts.append(t_el.text)

        elif child_tag == "f":
            # m:f → fraction: numerator/denominator
            num_el = child.find(f"{{{_M_NS}}}num")
            den_el = child.find(f"{{{_M_NS}}}den")
            num_latex = _convert_omml_to_latex_simplified(num_el) if num_el is not None else ""
            den_latex = _convert_omml_to_latex_simplified(den_el) if den_el is not None else ""
            if num_latex or den_latex:
                parts.append(f"\\frac{{{num_latex}}}{{{den_latex}}}")

        elif child_tag == "sSup":
            # m:sSup → superscript
            base_el = child.find(f"{{{_M_NS}}}e")
            sup_el = child.find(f"{{{_M_NS}}}sup")
            base_latex = _convert_omml_to_latex_simplified(base_el) if base_el is not None else ""
            sup_latex = _convert_omml_to_latex_simplified(sup_el) if sup_el is not None else ""
            parts.append(f"{{{base_latex}}}^{{{sup_latex}}}")

        elif child_tag == "sSub":
            # m:sSub → subscript
            base_el = child.find(f"{{{_M_NS}}}e")
            sub_el = child.find(f"{{{_M_NS}}}sub")
            base_latex = _convert_omml_to_latex_simplified(base_el) if base_el is not None else ""
            sub_latex = _convert_omml_to_latex_simplified(sub_el) if sub_el is not None else ""
            parts.append(f"{{{base_latex}}}_{{{sub_latex}}}")

        elif child_tag == "d":
            # m:d → delimiter (parentheses, brackets, etc.)
            begchr = ""
            endchr = ""
            dPr = child.find(f"{{{_M_NS}}}dPr")
            if dPr is not None:
                beg_el = dPr.find(f"{{{_M_NS}}}begChr")
                end_el = dPr.find(f"{{{_M_NS}}}endChr")
                if beg_el is not None:
                    begchr = beg_el.get(f"{{{_M_NS}}}val", "(")
                else:
                    begchr = "("
                if end_el is not None:
                    endchr = end_el.get(f"{{{_M_NS}}}val", ")")
                else:
                    endchr = ")"
            else:
                begchr = "("
                endchr = ")"

            e_parts = []
            for e_child in child.findall(f"{{{_M_NS}}}e"):
                e_latex = _convert_omml_to_latex_simplified(e_child)
                if e_latex:
                    e_parts.append(e_latex)
            inner = " ".join(e_parts)
            parts.append(f"{begchr}{inner}{endchr}")

    result = "".join(parts).strip()
    # Clean up excessive braces for readability
    result = re.sub(r"\{(\w)\}", r"\1", result)
    return result


def _extract_docx_table_rows(table_element) -> Dict[str, Any]:
    """Extract table rows with merged-cell handling (LightRAG style).

    Handles both horizontal (gridSpan) and vertical (vMerge) merges.
    Returns dict with "rows", "header_indices", "num_cols".
    """
    from xml.etree import ElementTree as ET

    qn_cache: Dict[str, str] = {}

    def qn(tag: str) -> str:
        if tag not in qn_cache:
            # Build qualified name for w: namespace
            if ":" in tag:
                prefix, local = tag.split(":", 1)
                ns_map = {"w": _W_NS, "wp": _WP_NS, "a": _A_NS, "r": _R_NS, "m": _M_NS}
                qn_cache[tag] = f"{{{ns_map.get(prefix, '')}}}{local}"
            else:
                qn_cache[tag] = f"{{{_W_NS}}}{tag}"
        return qn_cache[tag]

    tbl = table_element

    # Get column count from tblGrid
    tbl_grid = tbl.find(qn("w:tblGrid"))
    num_cols = 0
    if tbl_grid is not None:
        num_cols = len(tbl_grid.findall(qn("w:gridCol")))

    if num_cols == 0:
        return {"rows": [], "header_indices": [], "num_cols": 0}

    # Detect header rows
    header_indices: list[int] = []
    for idx, tr in enumerate(tbl.findall(qn("w:tr"))):
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None:
            tbl_header = trPr.find(qn("w:tblHeader"))
            if tbl_header is not None:
                header_indices.append(idx)

    grid: list[list[str]] = []
    vmerge_content: Dict[int, Dict[str, Any]] = {}  # {col: {text, para_id}}

    for tr in tbl.findall(qn("w:tr")):
        row_data = [""] * num_cols
        grid_col = 0

        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))

            # gridSpan (horizontal merge)
            grid_span = 1
            if tcPr is not None:
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is not None:
                    grid_span = int(gs.get(qn("w:val")))

            # vMerge (vertical merge)
            vmerge_elem = None
            vmerge_val = None
            if tcPr is not None:
                vmerge_elem = tcPr.find(qn("w:vMerge"))
                if vmerge_elem is not None:
                    vmerge_val = vmerge_elem.get(qn("w:val"))

            is_vmerge_restart = vmerge_elem is not None and vmerge_val == "restart"
            is_vmerge_continue = vmerge_elem is not None and vmerge_val in (None, "continue")
            is_normal_cell = vmerge_elem is None

            cell_text = ""

            if is_vmerge_restart or is_normal_cell:
                # Extract cell text from paragraphs
                para_texts: list[str] = []
                for para_el in tc.findall(qn("w:p")):
                    para_text = _extract_paragraph_content_omml(para_el, _DOCX_NS)
                    para_text = para_text.strip().replace("\x07", "")
                    if para_text:
                        para_texts.append(para_text)
                cell_text = "\n".join(para_texts)

                if is_vmerge_restart:
                    vmerge_content[grid_col] = {"text": cell_text}
                elif is_normal_cell and not cell_text and grid_col in vmerge_content:
                    cell_text = vmerge_content[grid_col]["text"]

            elif is_vmerge_continue:
                if grid_col in vmerge_content:
                    cell_text = vmerge_content[grid_col]["text"]

            if grid_col < num_cols:
                row_data[grid_col] = cell_text
            grid_col += grid_span

        grid.append(row_data)

    return {
        "rows": grid,
        "header_indices": header_indices,
        "num_cols": num_cols,
    }


def _extract_docx_drawing_placeholder(drawing_elem, relationships: Dict[str, _DocxRelationship]) -> str:
    """Build a placeholder string from a w:drawing element.

    Resolves a:blip → r:embed / r:link relationships.
    """
    doc_pr = drawing_elem.find(".//wp:docPr", _DOCX_NS)
    attrs = {
        "id": doc_pr.get("id", "") if doc_pr is not None else "",
        "name": doc_pr.get("name", "") if doc_pr is not None else "",
    }

    # Find blip
    for blip in drawing_elem.findall(".//a:blip", _DOCX_NS):
        rel_link = blip.get(f"{{{_R_NS}}}link")
        rel_embed = blip.get(f"{{{_R_NS}}}embed")

        rel_id = rel_link or rel_embed
        if rel_id and rel_id in relationships:
            rel = relationships[rel_id]
            if rel.image_format:
                attrs["format"] = rel.image_format
            if rel.target_mode.lower() == "external":
                attrs["path"] = rel.target
            elif rel.target:
                attrs["path"] = rel.target

    pieces = [f'{k}="{v}"' for k, v in attrs.items() if v]
    return f"<drawing {' '.join(pieces)} />"


def _parse_docx(source: Union[str, Path, bytes]) -> DocumentExtractionResult:
    """Parse DOCX using python-docx, extracting headings, tables, images, equations.

    LightRAG-style heading-level block splitting:
    - Each heading starts a new ContentBlock
    - Paragraphs accumulate under the current heading
    - Tables, images, equations produce dedicated blocks + side collections

    Args:
        source: File path string/Path, or raw DOCX bytes.

    Returns:
        DocumentExtractionResult with blocks, images, tables, equations.
    """
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Run: pip install python-docx")

    # --- open DOCX ---
    is_bytes = isinstance(source, bytes)
    if is_bytes:
        # python-docx can open from a file-like object
        doc = Document(io.BytesIO(source))
        source_label = "<bytes>"
        # For relationship/image extraction from bytes, we need a temp file or zip
        relationships: Dict[str, _DocxRelationship] = {}
        styles_outline: Dict[str, int] = {}
        # Write bytes to a temp file for zip access
        import tempfile

        tmp_path = tempfile.mktemp(suffix=".docx")
        Path(tmp_path).write_bytes(source)
        relationships = _load_docx_relationships(tmp_path)
        styles_outline = _parse_styles_outline_levels(tmp_path)
    else:
        docx_path = str(Path(source).resolve())
        try:
            doc = Document(docx_path)
        except PackageNotFoundError:
            raise ValueError(f"File is not a valid DOCX: {docx_path}")
        source_label = docx_path
        relationships = _load_docx_relationships(docx_path)
        styles_outline = _parse_styles_outline_levels(docx_path)

    result = DocumentExtractionResult(
        source_file=source_label,
        source_format="docx",
    )

    # --- iterate body elements ---
    body = doc._element.body
    block_idx = 0
    img_idx = 0
    tbl_idx = 0
    eq_idx = 0

    # Heading-level block splitting state
    current_heading = "Preface/Uncategorized"
    current_heading_level = 1
    heading_stack: Dict[int, str] = {}  # {level: heading_text}
    current_parent_headings: list[str] = []
    current_paragraphs: list[str] = []
    current_block_start_id = ""

    def _flush_block() -> None:
        nonlocal current_paragraphs, current_block_start_id, block_idx
        content = "\n".join(current_paragraphs).strip()
        if not content:
            current_paragraphs = []
            current_block_start_id = ""
            return

        parents = list(current_parent_headings)
        result.blocks.append(
            ContentBlock(
                block_id=f"blk_docx_{block_idx}",
                heading=current_heading,
                heading_level=current_heading_level,
                content=content,
                block_type="paragraph",
                source_id=current_block_start_id or f"blk_docx_{block_idx}",
            )
        )
        block_idx += 1
        current_paragraphs = []
        current_block_start_id = ""

    for element in body:
        tag = element.tag.split("}")[-1]

        if tag == "sectPr":
            continue

        if tag == "p":
            para_text = _extract_paragraph_content_omml(element, _DOCX_NS)
            para_text = para_text.strip()
            if not para_text:
                continue

            # Extract inline equations from paragraph text
            eq_matches = re.findall(r"<equation>(.*?)</equation>", para_text)
            for eq_latex in eq_matches:
                result.equations.append(
                    ExtractedEquation(
                        equation_id=f"eq_docx_{eq_idx}",
                        latex=eq_latex,
                        is_block=False,
                        source_id=f"eq_docx_{eq_idx}",
                    )
                )
                eq_idx += 1

            # Extract inline drawing placeholders → resolve to ExtractedImage
            drawing_matches = re.findall(r'<drawing\s+([^>]*?)\s*/>', para_text)
            for drawing_attrs_str in drawing_matches:
                attr_pairs = re.findall(r'(\w+)="([^"]*)"', drawing_attrs_str)
                attrs = dict(attr_pairs)
                rel_id = attrs.get("path", "")
                img_format = attrs.get("format", "jpeg")

                # Try to load embedded image bytes
                if not is_bytes:
                    rel = relationships.get(rel_id)
                    if rel and rel.part_name:
                        img_bytes = _extract_embedded_image_bytes(docx_path, rel.part_name)
                        if img_bytes:
                            compressed = _compress_image_bytes(img_bytes)
                            b64 = base64.b64encode(compressed).decode("ascii")
                            result.images.append(
                                ExtractedImage(
                                    image_id=f"img_docx_{img_idx}",
                                    base64_data=b64,
                                    format=img_format or "jpeg",
                                    source_id=rel_id,
                                    alt_text=attrs.get("name", ""),
                                )
                            )
                            img_idx += 1

            # Check heading level
            outline_level = _get_heading_level(element, styles_outline)
            para_id = _extract_para_id(element)

            if outline_level is not None and len(para_text) <= MAX_HEADING_LENGTH:
                # This is a heading
                level = outline_level + 1  # Convert 0-based to 1-based

                # Flush previous block
                _flush_block()

                # Strip any existing markdown heading prefix
                clean_heading = re.sub(r"^#{1,6}\s+", "", para_text)

                # Start new block with heading
                heading_prefix = "#" * min(level, 6)
                current_paragraphs = [f"{heading_prefix} {para_text}"]
                current_heading = clean_heading
                current_heading_level = level
                current_block_start_id = para_id or ""

                # Update heading stack
                heading_stack = {k: v for k, v in heading_stack.items() if k < level}
                heading_stack[level] = clean_heading
                current_parent_headings = [
                    heading_stack[lvl]
                    for lvl in sorted(heading_stack.keys())
                    if lvl < level
                ]
            else:
                # Regular paragraph
                if not current_block_start_id and para_id:
                    current_block_start_id = para_id
                current_paragraphs.append(para_text)

        elif tag == "tbl":
            # Table extraction with merged-cell handling
            table_data = _extract_docx_table_rows(element)
            table_rows = table_data["rows"]
            header_indices = table_data["header_indices"]

            # Skip empty tables
            if not table_rows or all(
                not (cell or "").strip() for row in table_rows for cell in row
            ):
                continue

            header_rows = None
            if header_indices:
                header_rows = [table_rows[idx] for idx in header_indices if idx < len(table_rows)]

            result.tables.append(
                ExtractedTable(
                    table_id=f"tbl_docx_{tbl_idx}",
                    rows=table_rows,
                    header_rows=header_rows,
                    num_rows=len(table_rows),
                    num_cols=table_data["num_cols"],
                    source_id=f"tbl_docx_{tbl_idx}",
                )
            )
            tbl_idx += 1

            # Add table placeholder to current block content
            table_json = json.dumps(table_rows, ensure_ascii=False)
            current_paragraphs.append(f"<table>{table_json}</table>")

    # Flush final block
    _flush_block()

    # --- extract embedded images from relationships ---
    if not is_bytes:
        for rel_id, rel in relationships.items():
            if rel.rel_type == _IMAGE_REL_TYPE and rel.target_mode.lower() != "external":
                if rel.part_name:
                    img_bytes = _extract_embedded_image_bytes(docx_path, rel.part_name)
                    if img_bytes:
                        compressed = _compress_image_bytes(img_bytes)
                        b64 = base64.b64encode(compressed).decode("ascii")
                        fmt = rel.image_format or "jpeg"
                        result.images.append(
                            ExtractedImage(
                                image_id=f"img_docx_rel_{img_idx}",
                                base64_data=b64,
                                format=fmt,
                                source_id=rel_id,
                            )
                        )
                        img_idx += 1

    # Clean up temp file if created
    if is_bytes and tmp_path:
        try:
            Path(tmp_path).unlink()
        except Exception:
            pass

    result.metadata["total_paragraphs"] = block_idx
    log.info(f"DOCX extraction complete: {result.summary()}")
    return result


# ============================================================================
# Markdown sub-parser — regex-based extraction (LightRAG style)
# ============================================================================

# Markdown token patterns
_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*$")
_MD_TRAILING_HASH_RE = re.compile(r"\s+#+\s*$")
_MD_FENCE_RE = re.compile(r"^(`{3,}|~{3,})(.*)$")
_MD_DELIMITER_ROW_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)*\|?\s*$")
_MD_DELIMITER_CELL_RE = re.compile(r"^:?-+:?$")
_MD_IMAGE_RE = re.compile(
    r'!\[(?P<alt>[^\]]*)\]\(\s*(?P<src><[^>]*>|[^)\s]+)(?:\s+"[^"]*")?\s*\)'
)
_MD_INLINE_EQ_RE = re.compile(r"\$((?:[^$\\]|\\.)+?)\$")  # $...$ inline math

_PREFACE_HEADING = "Preface/Uncategorized"


def _split_pipe_row(line: str) -> List[str]:
    """Split a pipe-table row into trimmed cells."""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [cell.strip() for cell in s.split("|")]


def _is_pipe_table_delimiter(header_line: str, delim_line: str) -> bool:
    """Check if delim_line is a valid GFM delimiter row for header_line."""
    if not _MD_DELIMITER_ROW_RE.match(delim_line):
        return False
    delim_cells = _split_pipe_row(delim_line)
    if not all(_MD_DELIMITER_CELL_RE.match(c) for c in delim_cells):
        return False
    return len(delim_cells) == len(_split_pipe_row(header_line))


def _consume_block_equation_md(lines: List[str], start: int) -> Tuple[int, str]:
    """Parse $$-delimited block equation. Returns (consumed, latex) or (0, "")."""
    first = lines[start].strip()
    inner = first[2:]

    # Single-line $$ ... $$ (but not just $$ at start and $$ at end of same line)
    if inner.rstrip().endswith("$$") and len(inner.rstrip()) >= 2:
        latex = inner.rstrip()[:-2].strip()
        return 1, latex

    # Multi-line
    body: list[str] = []
    if inner.strip():
        body.append(inner.strip())
    j = start + 1
    while j < len(lines):
        s = lines[j].strip()
        if s.endswith("$$"):
            tail = s[:-2].strip()
            if tail:
                body.append(tail)
            return (j - start + 1), "\n".join(body).strip()
        body.append(lines[j])
        j += 1
    return 0, ""


def _consume_pipe_table_md(lines: List[str], start: int) -> Tuple[int, List[List[str]], Optional[List[List[str]]]]:
    """Parse a GFM pipe table starting at lines[start]."""
    header = _split_pipe_row(lines[start])
    body_rows: list[list[str]] = []
    j = start + 2
    while j < len(lines):
        s = lines[j].strip()
        if not s or "|" not in s:
            break
        body_rows.append(_split_pipe_row(lines[j]))
        j += 1
    return (j - start), body_rows, [header] if header else None


def _resolve_md_image(src: str, md_dir: Optional[Path] = None) -> Optional[ExtractedImage]:
    """Resolve a markdown image reference to ExtractedImage.

    Handles:
    - Local file references (relative to markdown directory)
    - Absolute file paths
    - HTTP/HTTPS URLs (base64 fetch)
    - Data URIs (base64 embedded)
    """
    src = src.strip()
    if src.startswith("<") and src.endswith(">"):
        src = src[1:-1].strip()

    # Data URI: data:image/png;base64,...
    if src.lower().startswith("data:"):
        match = re.match(r"data:image/([\w+.-]+);base64,(.*)", src, re.DOTALL)
        if match:
            fmt = _normalize_image_format(match.group(1)) or "jpeg"
            b64_data = match.group(2)
            return ExtractedImage(
                image_id=f"img_md_data",
                base64_data=b64_data,
                format=fmt,
                source_id="data-uri",
            )
        return None

    # HTTP/HTTPS URL — fetch and base64 encode
    if src.startswith(("http://", "https://")):
        try:
            import urllib.request

            with urllib.request.urlopen(src, timeout=10) as resp:
                img_bytes = resp.read()
            compressed = _compress_image_bytes(img_bytes)
            b64 = base64.b64encode(compressed).decode("ascii")
            suffix = PurePosixPath(src).suffix.lstrip(".")
            fmt = _normalize_image_format(suffix) or "jpeg"
            return ExtractedImage(
                image_id=f"img_md_url",
                base64_data=b64,
                format=fmt,
                source_id=src,
            )
        except Exception as e:
            log.warning(f"Failed to fetch markdown image URL {src}: {e}")
            return None

    # Local file reference
    if md_dir is not None:
        local_path = md_dir / src
        if local_path.exists():
            try:
                img_bytes = local_path.read_bytes()
                compressed = _compress_image_bytes(img_bytes)
                b64 = base64.b64encode(compressed).decode("ascii")
                suffix = local_path.suffix.lstrip(".")
                fmt = _normalize_image_format(suffix) or "jpeg"
                return ExtractedImage(
                    image_id=f"img_md_local",
                    base64_data=b64,
                    format=fmt,
                    source_id=str(local_path),
                )
            except Exception as e:
                log.warning(f"Failed to read local markdown image {local_path}: {e}")

    return None


def _parse_markdown(source: Union[str, Path, bytes]) -> DocumentExtractionResult:
    """Parse Markdown using regex-based extraction (LightRAG style).

    Handles: ATX headings, GFM pipe tables, $$block equations, ![alt](src) images,
    inline $...$ equations, fenced code blocks.

    Args:
        source: File path string/Path, or raw Markdown bytes/str.

    Returns:
        DocumentExtractionResult with blocks, images, tables, equations.
    """
    # --- load markdown content ---
    md_dir: Optional[Path] = None
    if isinstance(source, bytes):
        text = source.decode("utf-8", errors="replace")
        source_label = "<bytes>"
    elif isinstance(source, Path):
        text = source.read_text(encoding="utf-8", errors="replace")
        md_dir = source.parent
        source_label = str(source.resolve())
    else:
        # source is a string — could be a file path or raw markdown content
        p = Path(source)
        if p.exists() and p.is_file():
            text = p.read_text(encoding="utf-8", errors="replace")
            md_dir = p.parent
            source_label = str(p.resolve())
        else:
            # Treat as raw markdown text content
            text = source
            source_label = "<text_input>"

    result = DocumentExtractionResult(
        source_file=source_label,
        source_format="md",
    )

    lines = text.splitlines()
    n = len(lines)

    block_idx = 0
    img_idx = 0
    tbl_idx = 0
    eq_idx = 0

    # Heading-level block splitting state
    heading_stack: list[tuple[int, str]] = []
    cur_heading = _PREFACE_HEADING
    cur_level = 0
    cur_parents: list[str] = []
    cur_lines: list[str] = []

    def _flush() -> None:
        nonlocal cur_lines, block_idx
        content = "\n".join(cur_lines).strip()
        if not content:
            cur_lines = []
            return
        result.blocks.append(
            ContentBlock(
                block_id=f"blk_md_{block_idx}",
                heading=cur_heading,
                heading_level=cur_level,
                content=content,
                block_type="paragraph",
                source_id=f"blk_md_{block_idx}",
            )
        )
        block_idx += 1
        cur_lines = []

    i = 0
    fence: Optional[Tuple[str, int]] = None  # (char, length)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # --- fenced code blocks: verbatim, suppress all detection ---
        fence_match = _MD_FENCE_RE.match(stripped)
        if fence is not None:
            cur_lines.append(line)
            if fence_match:
                ch = fence_match.group(1)[0]
                run = len(fence_match.group(1))
                if ch == fence[0] and run >= fence[1] and not fence_match.group(2):
                    fence = None
            i += 1
            continue
        if fence_match:
            fence = (fence_match.group(1)[0], len(fence_match.group(1)))
            cur_lines.append(line)
            i += 1
            continue

        # --- ATX heading ---
        heading_match = _MD_HEADING_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            raw = heading_match.group(2)
            clean = _MD_TRAILING_HASH_RE.sub("", raw).strip()

            _flush()
            heading_stack[:] = heading_stack[: max(level - 1, 0)]
            parents = [h for _, h in heading_stack if h]
            heading_stack.append((level, clean))
            cur_heading = clean
            cur_level = level
            cur_parents = parents
            cur_lines.append(f"{'#' * min(level, 6)} {raw}")
            i += 1
            continue

        # --- block equation ($$ ... $$) ---
        if stripped.startswith("$$"):
            consumed, latex = _consume_block_equation_md(lines, i)
            if consumed > 0:
                result.equations.append(
                    ExtractedEquation(
                        equation_id=f"eq_md_{eq_idx}",
                        latex=latex,
                        is_block=True,
                        source_id=f"eq_md_{eq_idx}",
                    )
                )
                eq_idx += 1
                cur_lines.append(f"<equation>{latex}</equation>")
                i += consumed
                continue

        # --- GFM pipe table ---
        if "|" in line and i + 1 < n and _is_pipe_table_delimiter(line, lines[i + 1]):
            consumed, body_rows, header_grid = _consume_pipe_table_md(lines, i)
            if consumed > 0:
                all_rows = (header_grid or []) + body_rows
                header_rows = header_grid if header_grid else None
                result.tables.append(
                    ExtractedTable(
                        table_id=f"tbl_md_{tbl_idx}",
                        rows=all_rows,
                        header_rows=header_rows,
                        num_rows=len(all_rows),
                        num_cols=max(len(r) for r in all_rows) if all_rows else 0,
                        source_id=f"tbl_md_{tbl_idx}",
                    )
                )
                tbl_idx += 1
                table_json = json.dumps(all_rows, ensure_ascii=False)
                cur_lines.append(f"<table>{table_json}</table>")
                i += consumed
                continue

        # --- inline images ---
        new_line = line
        for img_match in _MD_IMAGE_RE.finditer(line):
            alt = img_match.group("alt")
            src = img_match.group("src").strip()
            if src.startswith("<") and src.endswith(">"):
                src = src[1:-1].strip()

            img = _resolve_md_image(src, md_dir)
            if img is not None:
                img.image_id = f"img_md_{img_idx}"
                img.alt_text = alt
                result.images.append(img)
                img_idx += 1
                # Replace markdown image syntax with placeholder
                placeholder = f'<drawing id="{img.image_id}" name="{alt}" />'
                new_line = new_line.replace(img_match.group(0), placeholder)

        # --- inline equations $...$ ---
        for inline_match in _MD_INLINE_EQ_RE.finditer(new_line):
            latex = inline_match.group(1)
            result.equations.append(
                ExtractedEquation(
                    equation_id=f"eq_md_inline_{eq_idx}",
                    latex=latex,
                    is_block=False,
                    source_id=f"eq_md_inline_{eq_idx}",
                )
            )
            eq_idx += 1
            placeholder = f"<equation>{latex}</equation>"
            new_line = new_line.replace(inline_match.group(0), placeholder, 1)

        cur_lines.append(new_line)
        i += 1

    _flush()

    result.metadata["total_lines"] = n
    log.info(f"Markdown extraction complete: {result.summary()}")
    return result


# ============================================================================
# Plain text fallback sub-parser
# ============================================================================


def _parse_text(source: Union[str, Path, bytes]) -> DocumentExtractionResult:
    """Parse plain text as a single block (fallback parser).

    Args:
        source: File path string/Path, or raw text bytes/str.

    Returns:
        DocumentExtractionResult with one ContentBlock containing all text.
    """
    if isinstance(source, bytes):
        text = source.decode("utf-8", errors="replace")
        source_label = "<bytes>"
    elif isinstance(source, Path):
        text = source.read_text(encoding="utf-8", errors="replace")
        source_label = str(source.resolve())
    else:
        p = Path(source)
        if p.exists() and p.is_file():
            text = p.read_text(encoding="utf-8", errors="replace")
            source_label = str(p.resolve())
        else:
            text = source
            source_label = "<text_input>"

    result = DocumentExtractionResult(
        source_file=source_label,
        source_format="txt",
        metadata={"total_chars": len(text)},
    )

    # Attempt simple line-based splitting: double newline = paragraph break
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [text.strip()] if text.strip() else []

    for idx, para in enumerate(paragraphs):
        result.blocks.append(
            ContentBlock(
                block_id=f"blk_txt_{idx}",
                heading="",
                heading_level=0,
                content=para,
                block_type="paragraph",
                source_id=f"blk_txt_{idx}",
            )
        )

    log.info(f"Text extraction complete: {result.summary()}")
    return result


# ============================================================================
# Register built-in parsers
# ============================================================================

ParserRegistry.register(".pdf", _parse_pdf)
ParserRegistry.register(".docx", _parse_docx)
ParserRegistry.register(".md", _parse_markdown)
ParserRegistry.register(".txt", _parse_text)


# ============================================================================
# UnifiedDocumentParser — main operator class
# ============================================================================


class UnifiedDocumentParser:
    """Unified document parser operator for HugeGraph-AI.

    Routes by file suffix to format-specific sub-parsers, all producing
    the same DocumentExtractionResult. Follows the HG-AI operator protocol:
    ``run(context) -> context``.

    Usage::

        parser = UnifiedDocumentParser()
        ctx = parser.run({
            "document_path": "/path/to/report.pdf",
        })
        extraction = ctx["document_extraction"]
        for block in extraction.blocks:
            print(block.heading, block.content[:80])

        # Or from bytes:
        ctx = parser.run({
            "document_content": pdf_bytes,
            "document_format": "pdf",
        })

    Args:
        max_image_size_kb: Maximum image base64 size in KB (compression target).
        min_image_dim: Minimum image dimension to include (filter tiny images).
    """

    def __init__(
        self,
        max_image_size_kb: int = 512,
        min_image_dim: int = 50,
    ):
        self.max_image_size_kb = max_image_size_kb
        self.min_image_dim = min_image_dim

    def run(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Execute document parsing and add extraction result to context.

        Args:
            context: Must contain one of:
              - ``document_path``: str — file path to parse
              - ``document_content``: bytes — raw document bytes
              When ``document_content`` is used, ``document_format`` (e.g. "pdf")
              must also be provided.

        Returns:
            context dict with added ``"document_extraction"`` key containing
            a DocumentExtractionResult instance.

        Raises:
            ValueError: When no document source is provided or format is unknown.
        """
        doc_path = context.get("document_path")
        doc_content = context.get("document_content")
        doc_text = context.get("document_text")  # raw text string (for md/txt)
        doc_format = context.get("document_format")

        if doc_path is None and doc_content is None and doc_text is None:
            raise ValueError(
                "context must contain 'document_path', 'document_content', or 'document_text'"
            )

        # Determine source and suffix
        if doc_path is not None:
            suffix = Path(doc_path).suffix.lower()
            source = doc_path
        elif doc_text is not None:
            if doc_format is None:
                doc_format = "md"  # default text format
            suffix = f".{doc_format.lower()}"
            source = doc_text
        elif doc_content is not None:
            if doc_format is None:
                raise ValueError(
                    "document_format is required when using document_content"
                )
            suffix = f".{doc_format.lower()}"
            source = doc_content
        else:
            raise ValueError("No document source provided")

        # Resolve aliases
        resolved_suffix = ParserRegistry._aliases.get(suffix, suffix)
        parser_fn = ParserRegistry.get_parser(suffix)

        if parser_fn is None:
            raise ValueError(
                f"No parser registered for format '{suffix}'. "
                f"Supported formats: {ParserRegistry.supported_formats()}"
            )

        log.info(f"Routing document to parser: {resolved_suffix}")

        try:
            extraction = parser_fn(source)
        except Exception as e:
            log.error(f"Document parsing failed for {source}: {e}")
            # Return empty result on failure rather than crashing the pipeline
            source_label = str(source) if not isinstance(source, bytes) else "<bytes>"
            extraction = DocumentExtractionResult(
                source_file=source_label,
                source_format=resolved_suffix.lstrip("."),
                metadata={"error": str(e)},
            )

        context["document_extraction"] = extraction
        log.info(f"Extraction result: {extraction.summary()}")
        return context


# ============================================================================
# CLI entry point
# ============================================================================

if __name__ == "__main__":
    import sys

    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

    if len(sys.argv) < 2:
        print("Usage: python unified_document_parser.py <file_path> [format]")
        print("  file_path: path to PDF/DOCX/MD/TXT file")
        print("  format:    optional override (pdf/docx/md/txt)")
        print(f"  supported: {ParserRegistry.supported_formats()}")
        sys.exit(1)

    file_path = sys.argv[1]
    format_override = sys.argv[2] if len(sys.argv) > 2 else None

    ctx = {"document_path": file_path}
    if format_override:
        ctx["document_format"] = format_override

    parser = UnifiedDocumentParser()
    result_ctx = parser.run(ctx)
    extraction = result_ctx["document_extraction"]

    print(json.dumps(extraction.summary(), indent=2, ensure_ascii=False))

    print(f"\n=== Blocks ({extraction.total_blocks}) ===")
    for block in extraction.blocks[:10]:
        preview = block.content[:80].replace("\n", " ")
        print(f"  [{block.heading_level}] {block.heading}: {preview}...")

    if extraction.total_images > 0:
        print(f"\n=== Images ({extraction.total_images}) ===")
        for img in extraction.images[:5]:
            print(f"  {img.image_id}: {img.format} {img.size}, b64_len={len(img.base64_data)}")

    if extraction.total_tables > 0:
        print(f"\n=== Tables ({extraction.total_tables}) ===")
        for tbl in extraction.tables[:5]:
            print(f"  {tbl.table_id}: {tbl.num_rows}x{tbl.num_cols}, caption='{tbl.caption}'")

    if extraction.total_equations > 0:
        print(f"\n=== Equations ({extraction.total_equations}) ===")
        for eq in extraction.equations[:5]:
            print(f"  {eq.equation_id}: block={eq.is_block}, latex='{eq.latex[:40]}...'")
